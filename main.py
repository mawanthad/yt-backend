from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from googleapiclient.discovery import build
from docx import Document
import uuid
import os
import yt_dlp
import locale

locale.setlocale(locale.LC_ALL, '')
API_KEY = "YOUR_YOUTUBE_DATA_API_KEY"  # Replace this

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"]
)

FILES_DIR = "files"
os.makedirs(FILES_DIR, exist_ok=True)

class ScrapeRequest(BaseModel):
    url: str

def resolve_channel_id(input_url):
    youtube = build("youtube", "v3", developerKey=API_KEY)

    if "/channel/" in input_url:
        return input_url.split("/")[-1]

    identifier = input_url.strip('/').split('/')[-1].replace('@', '')

    request = youtube.search().list(
        part="snippet",
        q=identifier,
        type="channel",
        maxResults=1
    )
    response = request.execute()
    if not response['items']:
        return None
    return response['items'][0]['snippet']['channelId']

def get_videos(channel_id, max_results=10):
    youtube = build("youtube", "v3", developerKey=API_KEY)
    request = youtube.search().list(
        part="snippet",
        channelId=channel_id,
        maxResults=max_results,
        order="date"
    )
    response = request.execute()
    videos = []

    for item in response['items']:
        if item['id']['kind'] == 'youtube#video':
            video_id = item['id']['videoId']
            stats = youtube.videos().list(part="statistics", id=video_id).execute()
            views = stats['items'][0]['statistics'].get('viewCount', '0')
            videos.append({
                "video_id": video_id,
                "title": item['snippet']['title'],
                "date": item['snippet']['publishedAt'][:10],
                "views": locale.format_string("%d", int(views), grouping=True)
            })
    return videos

def fetch_transcript_yt_dlp(video_id):
    try:
        url = f"https://www.youtube.com/watch?v={video_id}"
        ydl_opts = {"skip_download": True, "writesubtitles": True, "writeautomaticsub": True}
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(url, download=False)
            subs = info.get("subtitles") or info.get("automatic_captions")
            if not subs or 'en' not in subs:
                return None
            sub_url = subs['en'][0]['url']
            from urllib.request import urlopen
            return urlopen(sub_url).read().decode("utf-8")
    except Exception as e:
        print("Transcript fetch failed:", e)
        return None

def create_docx(videos):
    doc = Document()
    for v in videos:
        doc.add_heading(v['title'], level=1)
        doc.add_paragraph(f"📅 Uploaded: {v['date']}")
        doc.add_paragraph(f"👁 Views: {v['views']}")
        doc.add_paragraph(v['transcript'])
        doc.add_page_break()

    filename = f"transcripts_{uuid.uuid4().hex[:8]}.docx"
    filepath = os.path.join(FILES_DIR, filename)
    doc.save(filepath)
    return filename

@app.post("/scrape")
async def scrape(data: ScrapeRequest):
    channel_id = resolve_channel_id(data.url)
    if not channel_id:
        return {"status": "error", "message": "Could not resolve channel."}

    videos = get_videos(channel_id)
    if not videos:
        return {"status": "error", "message": "No videos found."}

    transcripts = []
    for v in videos:
        transcript = fetch_transcript_yt_dlp(v['video_id'])
        if transcript:
            v['transcript'] = transcript
            transcripts.append(v)

    if not transcripts:
        return {"status": "no_transcripts", "message": "No transcripts found."}

    filename = create_docx(transcripts)
    return {
        "status": "success",
        "file": filename,
        "count": len(transcripts)
    }

@app.get("/files/{filename}")
def serve_file(filename: str):
    path = os.path.join(FILES_DIR, filename)
    if not os.path.exists(path):
        return {"message": "File not found"}
    return {"message": "File found"}
